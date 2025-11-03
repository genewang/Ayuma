# document_processor.py
import os
import logging
from typing import Dict, List, Optional
from pathlib import Path
import json
from datetime import datetime

# Import LangChain components
try:
    # Document loaders from langchain_community
    from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredMarkdownLoader
    from langchain_community.document_loaders import Docx2txtLoader, UnstructuredPowerPointLoader
    from langchain_community.document_loaders import UnstructuredExcelLoader, UnstructuredHTMLLoader

    # Core LangChain components
    try:
        from langchain.schema import Document
    except ImportError:
        from langchain_core.documents import Document

    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

    try:
        from langchain_chroma import Chroma
    except ImportError:
        try:
            from langchain.vectorstores import Chroma
        except ImportError:
            from langchain_community.vectorstores import Chroma

    try:
        from langchain.embeddings import HuggingFaceEmbeddings
    except ImportError:
        from langchain_huggingface import HuggingFaceEmbeddings

    # Force LangChain to be available
    LANGCHAIN_AVAILABLE = True
    logging.info("LangChain components successfully imported for document processing")
except ImportError as e:
    logging.error(f"LangChain import failed: {e}")
    # Still attempt to use fallback methods
    LANGCHAIN_AVAILABLE = False
    # Create fallback classes
    class Document:
        def __init__(self, page_content: str, metadata: dict = None):
            self.page_content = page_content
            self.metadata = metadata or {}
    RecursiveCharacterTextSplitter = None
    Chroma = None
    HuggingFaceEmbeddings = None

# Import fallback libraries for when LangChain is not available
try:
    import openpyxl
    OPENPYLIX_AVAILABLE = True
except ImportError:
    OPENPYLIX_AVAILABLE = False

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    import epub
    import ebooklib
    from bs4 import BeautifulSoup
    FALLBACK_AVAILABLE = True
except ImportError:
    FALLBACK_AVAILABLE = False

class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)

        # Initialize LangChain components
        self._initialize_langchain_components()

    def _initialize_langchain_components(self):
        """Initialize LangChain components"""
        if not LANGCHAIN_AVAILABLE:
            self.logger.warning("LangChain not available, using fallback document processing")
            self.embedding_model = None
            return

        if all([Document, RecursiveCharacterTextSplitter, Chroma, HuggingFaceEmbeddings]):
            try:
                # Initialize embedding model for metadata enhancement
                self.embedding_model = HuggingFaceEmbeddings(
                    model_name='sentence-transformers/all-MiniLM-L6-v2'
                )
                self.logger.info("Successfully initialized LangChain components")
            except Exception as e:
                self.logger.error(f"Failed to initialize LangChain components: {e}")
                self.embedding_model = None
        else:
            self.logger.warning("LangChain components not fully available, using fallback methods")
            self.embedding_model = None

    def _load_with_langchain(self, file_path: str) -> Optional[Document]:
        """Load a document using LangChain loaders"""
        if not LANGCHAIN_AVAILABLE:
            self.logger.warning(f"LangChain not available, cannot load {file_path}")
            return None
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                try:
                    loader = PyPDFLoader(file_path)
                    docs = loader.load()
                    if not docs or not docs[0].page_content.strip():
                        self.logger.warning(f"Empty content in PDF: {file_path}")
                        return None
                    return docs[0]
                except Exception as e:
                    self.logger.warning(f"Failed to load PDF with PyPDF: {e}")
                    # Try alternative PDF loader
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(file_path)
                        text = "\n\n".join(page.extract_text() for page in reader.pages)
                        if not text.strip():
                            self.logger.warning(f"No text extracted from PDF: {file_path}")
                            return None
                        return Document(page_content=text, metadata={"source": file_path})
                    except Exception as e2:
                        self.logger.warning(f"Failed to load PDF with fallback: {e2}")
                        return None
                        
            elif file_ext == '.docx':
                try:
                    # First try with docx2txt
                    loader = Docx2txtLoader(file_path)
                    docs = loader.load()
                    if docs and docs[0].page_content.strip():
                        return docs[0]
                except Exception as e:
                    self.logger.warning(f"Failed to load DOCX with docx2txt: {e}")
                
                # Fallback to python-docx
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs)
                    if not text.strip():
                        self.logger.warning(f"No text extracted from DOCX: {file_path}")
                        return None
                    return Document(page_content=text, metadata={"source": file_path})
                except Exception as e:
                    self.logger.warning(f"Failed to load DOCX with python-docx: {e}")
                    return None
                    
            # Handle other file types with simpler error handling
            loaders = {
                '.pptx': UnstructuredPowerPointLoader,
                '.xlsx': UnstructuredExcelLoader,
                '.xls': UnstructuredExcelLoader,
                '.html': UnstructuredHTMLLoader,
                '.htm': UnstructuredHTMLLoader,
                '.md': UnstructuredMarkdownLoader,
                '.txt': TextLoader
            }
            
            if file_ext in loaders:
                try:
                    loader = loaders[file_ext](file_path)
                    docs = loader.load()
                    if docs and docs[0].page_content.strip():
                        return docs[0]
                    self.logger.warning(f"Empty content in {file_ext} file: {file_path}")
                except Exception as e:
                    self.logger.warning(f"Failed to load {file_ext} with LangChain: {e}")
            else:
                self.logger.warning(f"Unsupported file type: {file_path}")
                
            return None
            
        except Exception as e:
            self.logger.error(f"Unexpected error in _load_with_langchain for {file_path}: {str(e)}", exc_info=True)
            return None
            
    def _load_with_fallback(self, file_path: str) -> Optional[Document]:
        """Load a document using fallback methods when LangChain is not available"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self.extract_text_from_excel(file_path)
            elif file_ext in ['.pptx', '.ppt']:
                return self.extract_text_from_powerpoint(file_path)
            elif file_ext == '.epub':
                return self.extract_text_from_epub(file_path)
            elif file_ext in ['.html', '.htm']:
                return self.extract_text_from_html(file_path)
            elif file_ext == '.txt':
                return self.extract_text_from_txt(file_path)
            else:
                self.logger.warning(f"Unsupported file type for fallback: {file_ext}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error in _load_with_fallback for {file_path}: {str(e)}", exc_info=True)
            return None
            
    def load_document_with_langchain(self, file_path: str) -> Optional[Document]:
        """Load document using appropriate LangChain document loader or fallback"""
        if not LANGCHAIN_AVAILABLE:
            self.logger.warning(f"LangChain not available, cannot load {file_path}")
            return None

        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext == '.pdf':
                try:
                    loader = PyPDFLoader(file_path)
                    docs = loader.load()
                    if not docs or not docs[0].page_content.strip():
                        self.logger.warning(f"Empty content in PDF: {file_path}")
                        return None
                    return docs[0]
                except Exception as e:
                    self.logger.warning(f"Failed to load PDF with PyPDF: {e}")
                    # Try alternative PDF loader
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(file_path)
                        text = "\n\n".join(page.extract_text() for page in reader.pages)
                        if not text.strip():
                            self.logger.warning(f"No text extracted from PDF: {file_path}")
                            return None
                        return Document(page_content=text, metadata={"source": file_path})
                    except Exception as e2:
                        self.logger.warning(f"Failed to load PDF with fallback: {e2}")
                        return None
                        
            elif file_ext == '.docx':
                try:
                    # First try with docx2txt
                    loader = Docx2txtLoader(file_path)
                    docs = loader.load()
                    if docs and docs[0].page_content.strip():
                        return docs[0]
                except Exception as e:
                    self.logger.warning(f"Failed to load DOCX with docx2txt: {e}")
                
                # Fallback to python-docx
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs)
                    if not text.strip():
                        self.logger.warning(f"No text extracted from DOCX: {file_path}")
                        return None
                    return Document(page_content=text, metadata={"source": file_path})
                except Exception as e:
                    self.logger.warning(f"Failed to load DOCX with python-docx: {e}")
                    return None
                    
            # Handle other file types with simpler error handling
            loaders = {
                '.pptx': UnstructuredPowerPointLoader,
                '.xlsx': UnstructuredExcelLoader,
                '.xls': UnstructuredExcelLoader,
                '.html': UnstructuredHTMLLoader,
                '.htm': UnstructuredHTMLLoader,
                '.md': UnstructuredMarkdownLoader,
                '.txt': TextLoader
            }
            
            if file_ext in loaders:
                try:
                    loader = loaders[file_ext](file_path)
                    docs = loader.load()
                    if docs and docs[0].page_content.strip():
                        return docs[0]
                    self.logger.warning(f"Empty content in {file_ext} file: {file_path}")
                except Exception as e:
                    self.logger.warning(f"Failed to load {file_ext} with LangChain: {e}")
            else:
                self.logger.warning(f"Unsupported file type: {file_path}")
                
            return None
            
        except Exception as e:
            self.logger.error(f"Unexpected error in _load_with_langchain for {file_path}: {str(e)}", exc_info=True)
            return None

    def load_document(self, file_path: str) -> Optional[Document]:
        """Load a document from file using LangChain if available, otherwise use fallback"""
        try:
            # Verify file exists and is not empty
            if not os.path.exists(file_path):
                self.logger.error(f"File not found: {file_path}")
                return None
                
            if os.path.getsize(file_path) == 0:
                self.logger.error(f"Empty file: {file_path}")
                return None
                
            # Try with LangChain first if available
            if LANGCHAIN_AVAILABLE:
                try:
                    return self._load_with_langchain(file_path)
                except Exception as e:
                    self.logger.warning(f"Failed to load {file_path} with LangChain: {e}")
                    self.logger.info("Falling back to direct parsing...")
            
            # Fallback to direct parsing
            return self._load_with_fallback(file_path)
            
        except Exception as e:
            self.logger.error(f"Unexpected error loading document {file_path}: {str(e)}", exc_info=True)
            return None

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text using LangChain document loaders or fallback methods"""
        if LANGCHAIN_AVAILABLE:
            try:
                doc = self.load_document(file_path)
                if doc:
                    return doc.page_content
                else:
                    return ""
            except Exception as e:
                self.logger.error(f"Error extracting text from {file_path}: {e}")
                return ""
        else:
            # Fallback to basic text reading for simple formats
            file_extension = Path(file_path).suffix.lower()
            if file_extension in ['.txt', '.md']:
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        return file.read()
                except Exception as e:
                    self.logger.error(f"Error reading text file {file_path}: {e}")
                    return ""
            else:
                self.logger.warning(f"Cannot process {file_extension} files without LangChain")
                return ""

    def extract_metadata_from_file(self, file_path: str) -> Dict:
        """Extract metadata from file using LangChain and file system info"""
        try:
            stat = os.stat(file_path)
            file_name = Path(file_path).name
            file_extension = Path(file_path).suffix.lower()

            metadata = {
                "file_name": file_name,
                "file_path": file_path,
                "file_extension": file_extension,
                "file_size": stat.st_size,
                "creation_date": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modification_date": datetime.fromtimestamp(stat.st_mtime).isoformat()
            }

            # Extract document-specific metadata using LangChain if available
            if LANGCHAIN_AVAILABLE:
                doc = self.load_document_with_langchain(file_path)
                if doc and hasattr(doc, 'metadata'):
                    # Merge LangChain metadata with our metadata
                    metadata.update(doc.metadata)

                # Add custom processing metadata
                metadata.update({
                    "processed_with": "langchain",
                    "loader_type": self._get_loader_type(file_extension),
                    "extraction_method": "langchain_document_loader"
                })
            else:
                # Fallback metadata without LangChain
                metadata.update({
                    "processed_with": "fallback",
                    "loader_type": "TextLoader" if file_extension in ['.txt', '.md'] else "Unknown",
                    "extraction_method": "basic_text_extraction"
                })

            return metadata

        except Exception as e:
            self.logger.error(f"Error extracting metadata from {file_path}: {e}")
            return {}

    def _get_loader_type(self, file_extension: str) -> str:
        """Get the LangChain loader type for the file extension"""
        loader_mapping = {
            '.pdf': 'PyPDFLoader',
            '.txt': 'TextLoader',
            '.md': 'UnstructuredMarkdownLoader',
            '.docx': 'Docx2txtLoader',
            '.doc': 'Docx2txtLoader',
            '.pptx': 'UnstructuredPowerPointLoader',
            '.ppt': 'UnstructuredPowerPointLoader',
            '.xlsx': 'UnstructuredExcelLoader',
            '.xls': 'UnstructuredExcelLoader',
            '.html': 'UnstructuredHTMLLoader',
            '.htm': 'UnstructuredHTMLLoader'
        }
        return loader_mapping.get(file_extension, 'UnknownLoader')

    def process_directory(self, directory_path: str, file_extensions: List[str] = None) -> List[Dict]:
        """Process all documents in a directory using LangChain loaders or fallbacks"""
        if file_extensions is None:
            # Support all LangChain-compatible formats
            file_extensions = ['.pdf', '.txt', '.md', '.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls', '.html', '.htm']

        documents = []
        directory = Path(directory_path)

        if not directory.exists():
            self.logger.error(f"Directory does not exist: {directory_path}")
            return documents

        # Find all files with specified extensions
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_extensions:
                try:
                    # Use LangChain to load the document if available
                    if LANGCHAIN_AVAILABLE:
                        doc = self.load_document_with_langchain(str(file_path))
                        if doc and doc.page_content.strip():
                            # Extract metadata using LangChain
                            metadata = self.extract_metadata_from_file(str(file_path))

                            # Create document structure for RAG system with LangChain metadata
                            document = {
                                "id": f"langchain_{file_path.stem}_{hash(file_path.name + str(metadata.get('modification_date', ''))) % 1000000}",
                                "content": doc.page_content,
                                "source": str(file_path),
                                "file_name": file_path.name,
                                "file_type": file_path.suffix,
                                "metadata": metadata,
                                "institution": self._extract_institution_from_content(doc.page_content) or "Unknown",
                                "evidence_level": self._extract_evidence_level_from_content(doc.page_content),
                                "document_type": self._extract_document_type_from_content(doc.page_content),
                                "cancer_type": self._extract_cancer_type_from_filename(file_path.name),
                                "publication_date": metadata.get("creation_date_pdf", metadata.get("creation_date", ""))[:10],
                                "langchain_metadata": doc.metadata if hasattr(doc, 'metadata') else {}
                            }

                            documents.append(document)
                            self.logger.info(f"Processed with LangChain: {file_path.name}")
                    else:
                        # Fallback processing for simple text files
                        text_content = self.extract_text_from_file(str(file_path))
                        if text_content.strip():
                            metadata = self.extract_metadata_from_file(str(file_path))

                            document = {
                                "id": f"fallback_{file_path.stem}_{hash(file_path.name + str(metadata.get('modification_date', ''))) % 1000000}",
                                "content": text_content,
                                "source": str(file_path),
                                "file_name": file_path.name,
                                "file_type": file_path.suffix,
                                "metadata": metadata,
                                "institution": self._extract_institution_from_content(text_content) or "Unknown",
                                "evidence_level": self._extract_evidence_level_from_content(text_content),
                                "document_type": self._extract_document_type_from_content(text_content),
                                "cancer_type": self._extract_cancer_type_from_filename(file_path.name),
                                "publication_date": metadata.get("creation_date_pdf", metadata.get("creation_date", ""))[:10],
                                "langchain_metadata": {}
                            }

                            documents.append(document)
                            self.logger.info(f"Processed with fallback: {file_path.name}")

                except Exception as e:
                    self.logger.error(f"Error processing {file_path}: {e}")

        processing_method = "LangChain" if LANGCHAIN_AVAILABLE else "Fallback"
        self.logger.info(f"Successfully processed {len(documents)} documents from {directory_path} using {processing_method}")
        return documents

    def _extract_institution_from_content(self, content: str) -> Optional[str]:
        """Extract institution/organization from document content"""
        content_lower = content.lower()

        # Look for common medical institutions and organizations
        institutions = {
            'national comprehensive cancer network': 'NCCN',
            'american society of clinical oncology': 'ASCO',
            'european society for medical oncology': 'ESMO',
            'national cancer institute': 'NCI',
            'american cancer society': 'ACS',
            'world health organization': 'WHO',
            'food and drug administration': 'FDA',
            'national institutes of health': 'NIH',
            'european medicines agency': 'EMA',
            'cochrane': 'COCHRANE',
            'uptodate': 'UPTODATE',
            'clinicaltrials.gov': 'CLINICALTRIALS',
            'national foundation for cancer research': 'NFCR',
            'nfcr': 'NFCR'
        }

        for full_name, abbreviation in institutions.items():
            if full_name in content_lower:
                return abbreviation

        return None

    def _extract_cancer_type_from_filename(self, filename: str) -> Optional[str]:
        """Extract cancer type from filename"""
        filename_lower = filename.lower()

        # Common cancer types
        cancer_types = [
            'breast cancer', 'lung cancer', 'prostate cancer', 'colorectal cancer',
            'melanoma', 'leukemia', 'lymphoma', 'pancreatic cancer', 'ovarian cancer',
            'bladder cancer', 'kidney cancer', 'liver cancer', 'stomach cancer',
            'esophageal cancer', 'thyroid cancer', 'brain cancer', 'cervical cancer',
            'endometrial cancer', 'testicular cancer', 'bone cancer', 'sarcoma',
            'myeloma', 'myelodysplastic', 'waldenstrom'
        ]

        for cancer_type in cancer_types:
            if cancer_type.replace(' ', '') in filename_lower.replace(' ', '').replace('_', '').replace('-', ''):
                return cancer_type

        return None

    def _extract_evidence_level_from_content(self, content: str) -> str:
        """Extract evidence level from document content"""
        content_lower = content.lower()

        # Evidence level indicators
        if any(term in content_lower for term in ['systematic review', 'meta-analysis', 'meta analysis']):
            return 'systematic-review'
        elif any(term in content_lower for term in ['randomized controlled trial', 'rct', 'randomized trial']):
            return 'randomized-controlled-trial'
        elif any(term in content_lower for term in ['cohort study', 'prospective study']):
            return 'cohort-study'
        elif any(term in content_lower for term in ['case control', 'case-control study']):
            return 'case-control'
        elif any(term in content_lower for term in ['guideline', 'consensus', 'expert opinion']):
            return 'expert-opinion'
        else:
            return 'unknown'

    def _extract_document_type_from_content(self, content: str) -> str:
        """Extract document type from content"""
        content_lower = content.lower()

        if any(term in content_lower for term in ['clinical trial', 'trial protocol', 'study protocol']):
            return 'clinical-trial'
        elif any(term in content_lower for term in ['review', 'literature review', 'systematic review']):
            return 'review'
        elif any(term in content_lower for term in ['guideline', 'recommendation', 'consensus']):
            return 'guideline'
        elif any(term in content_lower for term in ['case study', 'case report']):
            return 'case-study'
        else:
            return 'general'

    def save_processed_documents(self, documents: List[Dict], output_file: str = "processed_langchain_documents.json"):
        """Save processed documents to JSON file for backup or further processing"""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(documents, f, indent=2, ensure_ascii=False)
            self.logger.info(f"Saved {len(documents)} processed documents to {output_file}")
        except Exception as e:
            self.logger.error(f"Error saving processed documents: {e}")

    def load_processed_documents(self, input_file: str = "processed_langchain_documents.json") -> List[Dict]:
        """Load previously processed documents from JSON file"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                documents = json.load(f)
            self.logger.info(f"Loaded {len(documents)} processed documents from {input_file}")
            return documents
        except Exception as e:
            self.logger.error(f"Error loading processed documents: {e}")
            return []

    # Keep backward compatibility methods for existing functionality
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_excel(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_powerpoint(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_epub(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_html(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)

    def extract_text_from_txt(self, file_path: str) -> str:
        """Backward compatibility method"""
        return self.extract_text_from_file(file_path)
